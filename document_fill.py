from docx import Document
from docx2pdf import convert
from utils.map_level import map_level
from docx.shared import RGBColor



from time import sleep
from csv_extraction import participant_dict, company_dict, country_dict, most_frequent_company
from definitions import definitions

# Load the Word document
doc_path = r"C:\Users\GOOD\Downloads\AIR.docx"  # Replace with your actual file path
doc = Document(doc_path)

print(company_dict)

date = '16/02/2025' # Initialize data variable with current date


summary_table = doc.tables[0]
'''for each in summary_table.rows:
    print(len(summary_table.rows))
    row_data = [cell.text for cell in each.cells]  # Extract text from each cell
    print(row_data)'''

detail_table1 = doc.tables[1]


detail_table2 = doc.tables[1]
for each in detail_table2.rows:
    print(len(detail_table2.rows))
    row_data = [cell.text for cell in each.cells]  # Extract text from each cell
    print(row_data)
#print(doc.tables)


for each,details in participant_dict.items():

    print(f'Name ----------> {each}')
    #print(f'Details -------> {details['Total Score']}')
    #print(f'Level -------> {[map_level(details['Total Score'])]}')
    #print(f'Definition ------> {definitions["Total Score"][map_level(details["Total Score"])]}')

# ---------------------- Modifying Table 1 (Summary Table) ----------------------

    # Row 1 is 'Name' and 'Date' 
    row0 = summary_table.rows[0] 
    row0.cells[1].text = each
    row0.cells[3].text = date   
                
    row_data = [cell.text.strip() for cell in row0.cells]  # Extract text from each cell
    print(row_data)


    # Row 2 is 'Company' and 'Definition'
    row1 = summary_table.rows[1] 
    row1.cells[1].text = details['Company']
    row1.cells[3].text = definitions['Total Score'][map_level(details['Total Score'])]
    for paragraph in row0.cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    row_data1 = [cell.text.strip() for cell in row1.cells]  # Extract text from each cell
    print(row_data1)

    #Row 3 is 'Country' and 'Total Score'
    row2 = summary_table.rows[2] 
    row2.cells[1].text = details['Country']
    row2.cells[2].text = str(details['Total Score'])
    for paragraph in row0.cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
                
    row_data2 = [cell.text.strip() for cell in row2.cells]  # Extract text from each cell
    print(row_data2)

    #Row 4 is 'Company' and 'Total Score'
    row3 = summary_table.rows[3] 
    row3.cells[1].text = details['Department']
    row3.cells[2].text = str(details['Total Score'])
    for paragraph in row0.cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
    row_data3 = [cell.text.strip() for cell in row3.cells]  # Extract text from each cell
    print(row_data3)

    #Row 5 is 'Company' and 'Total Score'
    row4 = summary_table.rows[4] 
    row3.cells[1].text = details['Position']
    row3.cells[2].text = map_level(details['Total Score'])
    for paragraph in row0.cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    row_data4 = [cell.text.strip() for cell in row4.cells]  # Extract text from each cell
    print(row_data4)



# ---------------------- Modifying Table 2 (Details Table) ----------------------

 
    print('--------------- Table 2 ---------------')

    row1 = detail_table1.rows[1]
    row1.cells[0].text = str(company_dict[most_frequent_company]['Technical Foundation'])
    row1.cells[1].text = str(company_dict[most_frequent_company]['Personal Readiness'])
    row1.cells[2].text = str(company_dict[most_frequent_company]['External Awareness'])
    row1.cells[3].text = str(company_dict[most_frequent_company]['Individual Score'])
    for paragraph in row0.cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)


    row_data1 = [cell.text.strip() for cell in row1.cells]  # Extract text from each cell
    print(row_data1)

    
    row2 = detail_table1.rows[2]
    row2.cells[0].text = str(details['Technical Foundation'])
    row2.cells[1].text = str(details['Personal Readiness'])
    row2.cells[2].text = str(details['External Awareness'])
    row2.cells[3].text = str(details['Individual Score'])
    for paragraph in row0.cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    row_data2 = [cell.text.strip() for cell in row2.cells]  # Extract text from each cell
    print(row_data2)

    row3 = detail_table1.rows[3]
    row3.cells[0].text = map_level(details['Technical Foundation'])
    row3.cells[1].text = map_level(details['Personal Readiness'])
    row3.cells[2].text = map_level(details['External Awareness'])
    row3.cells[3].text = map_level(details['Individual Score'])
    for paragraph in row0.cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    row_data3 = [cell.text.strip() for cell in row3.cells]  # Extract text from each cell
    print(row_data3)

    row4 = detail_table1.rows[4]
    row4.cells[0].text = definitions['Technical Foundation'][map_level(details['Technical Foundation'])]
    row4.cells[1].text = definitions['Personal Readiness'][map_level(details['Personal Readiness'])]
    row4.cells[2].text = definitions['External Awareness'][map_level(details['External Awareness'])]
    row4.cells[3].text = definitions['Individual Score'][map_level(details['Individual Score'])]
    for paragraph in row0.cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    row_data4 = [cell.text.strip() for cell in row4.cells]  # Extract text from each cell
    print(row_data4)



    #---------- Table 3 ----------------------

    print('--------------- Table 3 ---------------')


    row1 = detail_table2.rows[1]
    row1.cells[0].text = str(company_dict[most_frequent_company]['Process Integration'])
    row1.cells[1].text = str(company_dict[most_frequent_company]['Department Integration'])
    row1.cells[2].text = str(company_dict[most_frequent_company]['Implementation Impact'])
    row1.cells[3].text = str(company_dict[most_frequent_company]['Integrated Score'])
    for paragraph in row0.cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)


    row_data1 = [cell.text.strip() for cell in row1.cells]  # Extract text from each cell
    print(row_data1)

    
    row2 = detail_table2.rows[2]
    row2.cells[0].text = str(details['Process Integration'])
    row2.cells[1].text = str(details['Department Integration'])
    row2.cells[2].text = str(details['Implementation Impact'])
    row2.cells[3].text = str(details['Integrated Score'])
    for paragraph in row0.cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    row_data2 = [cell.text.strip() for cell in row2.cells]  # Extract text from each cell
    print(row_data2)

    row3 = detail_table2.rows[3]
    row3.cells[0].text = map_level(details['Process Integration'])
    row3.cells[1].text = map_level(details['Department Integration'])
    row3.cells[2].text = map_level(details['Implementation Impact'])
    row3.cells[3].text = map_level(details['Integrated Score'])
    for paragraph in row0.cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    row_data3 = [cell.text.strip() for cell in row3.cells]  # Extract text from each cell
    print(row_data3)

    row4 = detail_table2.rows[4]
    row4.cells[0].text = definitions['Process Integration'][map_level(details['Process Integration'])]
    row4.cells[1].text = definitions['Department Integration'][map_level(details['Department Integration'])]
    row4.cells[2].text = definitions['Implementation Impact'][map_level(details['Implementation Impact'])]
    row4.cells[3].text = definitions['Integrated Score'][map_level(details['Integrated Score'])]
    for paragraph in row0.cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    row_data4 = [cell.text.strip() for cell in row4.cells]  # Extract text from each cell
    print(row_data4)

    # Save the updated document as a DOCX file (using python-docx)
    updated_doc_path = fr"C:\Users\Yeshwanth\Downloads\Test\{each}.docx"
    pdf_file = fr"C:\Users\Yeshwanth\Downloads\Test\{each}.pdf"

    doc.save(updated_doc_path)

    convert(updated_doc_path,pdf_file)